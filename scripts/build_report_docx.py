"""
Render REPORT.md to a Word document for submission.

    python scripts/build_report_docx.py                  # -> REPORT.docx
    python scripts/build_report_docx.py --out 2024XY.docx

Pandoc is not assumed to be installed, and asking a marker to accept a Markdown
file is not an option -- the assignment asks for a Word document. This handles
the subset of Markdown the report actually uses: headings, paragraphs, bullet
and numbered lists, tables, block quotes, fenced code, bold/italic/inline-code
spans, and horizontal rules.

Mermaid blocks cannot be rendered here, so they are emitted as a captioned code
block with a visible note -- paste an exported PNG over it. Silently dropping a
diagram from a graded document would be worse than an obvious placeholder.
"""

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent

# Usable height of a Letter page after default 1in margins, less room for a
# caption. Figures taller than this get scaled down rather than split.
MAX_FIGURE_IN = 8.0

# `code`, **bold**, *italic* -- ordered so the code pattern wins first and its
# contents are not re-scanned for emphasis markers.
_SPAN = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|(?<!\*)\*(?!\*)[^*]+\*(?!\*))")


def add_runs(paragraph, text: str) -> None:
    """Write text into a paragraph, honouring inline markup."""
    # Links: keep the label, drop the target -- a printed docx cannot be clicked
    # and the raw URL is noise in a report.
    text = re.sub(r"!?\[([^\]]+)\]\([^)]+\)", r"\1", text)

    for piece in _SPAN.split(text):
        if not piece:
            continue
        run = paragraph.add_run(piece)
        if piece.startswith("`") and piece.endswith("`"):
            run.text = piece[1:-1]
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0xC0, 0x30, 0x30)
        elif piece.startswith("**"):
            run.text = piece[2:-2]
            run.bold = True
        elif piece.startswith("*"):
            run.text = piece[1:-1]
            run.italic = True


def split_row(line: str):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def is_separator(line: str) -> bool:
    return bool(re.fullmatch(r"\|[\s:|-]+\|", line.strip()))


def build(md_path: Path, out_path: Path, student_id: str = "", name: str = "") -> None:
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    source = md_path.read_text(encoding="utf-8")

    # REPORT.md is committed to a public repository, so it carries placeholders
    # rather than a student's name and ID. They are substituted here, into the
    # document that actually gets submitted.
    source = source.replace("{{BITS_ID}}", student_id or "<BITS ID>")
    source = source.replace("{{NAME}}", name or "<Name>")
    # Drop HTML comments -- build instructions for the repo, noise in the report.
    source = re.sub(r"<!--.*?-->>?", "", source, flags=re.DOTALL)

    lines = source.splitlines()
    i, n = 0, len(lines)
    tables = code_blocks = images = 0
    missing = []
    tall = []

    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Fenced code ------------------------------------------------------
        if stripped.startswith("```"):
            language = stripped[3:].strip().lower()
            i += 1
            body = []
            while i < n and not lines[i].strip().startswith("```"):
                body.append(lines[i])
                i += 1
            i += 1  # closing fence

            if language == "mermaid":
                # If the rendered PNG was already inserted just above, this is
                # the source listing for it, not a missing-figure placeholder.
                rendered = images > 0
                note = doc.add_paragraph()
                run = note.add_run(
                    "Mermaid source for the diagram above:" if rendered
                    else "[Architecture diagram — paste the exported PNG here. Source below.]"
                )
                run.italic = rendered
                run.bold = not rendered
                run.font.size = Pt(9)
                if not rendered:
                    run.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)

            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Pt(18)
            run = para.add_run("\n".join(body))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            code_blocks += 1
            continue

        # Tables -----------------------------------------------------------
        if stripped.startswith("|") and i + 1 < n and is_separator(lines[i + 1]):
            header = split_row(stripped)
            i += 2
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                rows.append(split_row(lines[i]))
                i += 1

            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Table Grid"
            for cell, text in zip(table.rows[0].cells, header):
                cell.paragraphs[0].text = ""
                add_runs(cell.paragraphs[0], text)
                for run in cell.paragraphs[0].runs:
                    run.bold = True

            for row in rows:
                cells = table.add_row().cells
                # A ragged row would raise on zip-to-cells; pad instead of dying
                # halfway through writing the document.
                for cell, text in zip(cells, row + [""] * (len(header) - len(row))):
                    cell.paragraphs[0].text = ""
                    add_runs(cell.paragraphs[0], text)
            doc.add_paragraph()
            tables += 1
            continue

        # Images -----------------------------------------------------------
        image = re.fullmatch(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if image:
            caption, src = image.group(1), image.group(2)
            path = (md_path.parent / src).resolve()
            if path.exists():
                # 6.2in fits the default page's text width; letting python-docx
                # use the PNG's native size overflows the margin on a 1500px
                # screenshot and silently crops it in print.
                picture = doc.add_picture(str(path), width=Inches(6.2))
                # A full-page capture of a long JSON response can come out taller
                # than the page, splitting one figure across two sheets. Cap the
                # height and let the width shrink to keep it on a single page.
                if picture.height > Inches(MAX_FIGURE_IN):
                    picture.width = int(picture.width * Inches(MAX_FIGURE_IN) / picture.height)
                    picture.height = Inches(MAX_FIGURE_IN)
                    tall.append(f"{src} ({picture.width / 914400:.1f}in wide after scaling)")
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                warn = doc.add_paragraph()
                run = warn.add_run(f"[MISSING IMAGE: {src}]")
                run.bold = True
                run.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)
                missing.append(src)
            if caption:
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run(caption)
                run.italic = True
                run.font.size = Pt(9)
            images += 1
            i += 1
            continue

        # Headings ---------------------------------------------------------
        heading = re.match(r"(#{1,6})\s+(.*)", stripped)
        if heading:
            level = len(heading.group(1))
            para = doc.add_heading("", level=min(level, 4))
            add_runs(para, heading.group(2))
            i += 1
            continue

        # Horizontal rule ---------------------------------------------------
        if re.fullmatch(r"-{3,}|\*{3,}|_{3,}", stripped):
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.add_run("• • •").font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            i += 1
            continue

        # Block quote -------------------------------------------------------
        if stripped.startswith(">"):
            body = []
            while i < n and lines[i].strip().startswith(">"):
                body.append(lines[i].strip().lstrip(">").strip())
                i += 1
            para = doc.add_paragraph(style="Intense Quote")
            add_runs(para, " ".join(b for b in body if b))
            continue

        # Lists -------------------------------------------------------------
        bullet = re.match(r"[-*+]\s+(.*)", stripped)
        number = re.match(r"\d+[.)]\s+(.*)", stripped)
        if bullet or number:
            style = "List Bullet" if bullet else "List Number"
            text = (bullet or number).group(1)
            # Continuation lines of the same item are indented, not new items.
            i += 1
            while i < n and lines[i].startswith("  ") and lines[i].strip() \
                    and not re.match(r"[-*+]\s|\d+[.)]\s", lines[i].strip()):
                text += " " + lines[i].strip()
                i += 1
            para = doc.add_paragraph(style=style)
            add_runs(para, text)
            continue

        # Paragraph ---------------------------------------------------------
        body = [stripped]
        i += 1
        while i < n and lines[i].strip() and not re.match(
            r"#{1,6}\s|[-*+]\s|\d+[.)]\s|\||>|```|-{3,}$", lines[i].strip()
        ):
            body.append(lines[i].strip())
            i += 1
        para = doc.add_paragraph()
        add_runs(para, " ".join(body))

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"wrote {out_path}")
    print(f"  {len(doc.paragraphs)} paragraphs, {tables} tables, {code_blocks} code blocks, {images} images")
    if missing:
        print(f"  {len(missing)} MISSING image(s): {missing}")
    for note in tall:
        print(f"  scaled to fit one page: {note}")

    # Scan the substituted text, not the file on disk: the placeholders in
    # REPORT.md are meant to be there, and warning about ones that were just
    # filled in trains you to ignore the warning.
    remaining = [
        ln for ln in source.splitlines()
        if any(m in ln for m in ("<fill>", "<pending>", "{{BITS_ID}}", "{{NAME}}"))
    ]
    if remaining:
        print(f"\n  {len(remaining)} placeholder(s) still in the document:")
        for ln in remaining:
            print(f"    {ln.strip()[:100]}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Render REPORT.md to .docx")
    parser.add_argument("--md", default=str(ROOT / "REPORT.md"))
    parser.add_argument("--out", default=str(ROOT / "REPORT.docx"))
    parser.add_argument("--id", dest="student_id", default="",
                        help="BITS ID, substituted for {{BITS_ID}}")
    parser.add_argument("--name", default="", help="student name, substituted for {{NAME}}")
    args = parser.parse_args()
    build(Path(args.md), Path(args.out), args.student_id, args.name)
